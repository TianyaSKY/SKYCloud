import logging
import tempfile
from pathlib import Path

import cv2
import requests
from docx import Document
from openai import OpenAI

from .file_util import convert_office_to_pdf, convert_pdf_to_images, extract_video_frames

# 配置日志
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# 支持的文本扩展名
text_exts = {'.md', '.cpp', '.h', '.c', '.py', '.txt', '.java', '.go', '.rs', '.js', '.ts', '.html', '.css',
             '.json', '.yaml', '.yml', '.sql', '.sh', '.bat'}


def _get_visual_urls(local_path: str) -> list:
    """
    将不同类型的文件转换为可供多模态模型识别的图片 URL 列表。
    """
    path = Path(local_path)
    ext = path.suffix.lower()
    image_urls = []

    # 定义支持的格式
    doc_exts = {'.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.pdf'}
    video_exts = {'.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv'}
    img_exts = {'.jpg', '.jpeg', '.png', '.bmp', '.webp'}
    radio_exts = {'.mp3', '.wav', '.flac', '.ogg'}
    
    if ext in radio_exts:
        raise ValueError('暂不支持音频文件的解析！！！')
        
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            if ext in doc_exts:
                target_pdf = local_path if ext == '.pdf' else convert_office_to_pdf(local_path, tmpdir)
                images = convert_pdf_to_images(target_pdf, tmpdir, max_pages=20)
                for img_path in images:
                    upload = upload_file(img_path)
                    if upload: image_urls.append(upload)

            elif ext in video_exts:
                frames = extract_video_frames(local_path, tmpdir, frame_count=20)
                for img_path in frames:
                    upload = upload_file(img_path)
                    if upload: image_urls.append(upload)
            elif ext in img_exts:
                # 文件直接上传
                image_urls.append(upload_file(local_path))
            elif ext in text_exts:
                return []  # 无需处理
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        except Exception as e:
            logger.error(f"Error processing visual content for {local_path}: {e}")
            raise e

    return image_urls


def desc_file(local_path: str, api: str, key: str, model: str) -> str:
    """
    生成文件的多模态描述。
    """
    image_urls = _get_visual_urls(local_path)
    if not image_urls:
        logger.info(f"No visual content for {local_path}")

    client = OpenAI(api_key=key, base_url=api, timeout=120)
    
    # 强制要求最终输出为英文
    prompt = (
        "You are a professional file analysis assistant. Please carefully observe the provided file content "
        "(which may be document screenshots, video keyframes, or images) and generate an accurate, professional, "
        "and concise description. The description should include: core theme, main content summary, and key information points. "
        "Output the description text directly in ENGLISH without any preamble or explanatory text. Do not use Markdown syntax. "
        "No need to describe the sequence of images."
    )

    content = [{"type": "text", "text": prompt}]
    for url in image_urls:
        content.append({"type": "image_url", "image_url": {"url": url}})
    content.append({"type": "text", "text": format_data(local_path)})  # 加入文本信息
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": content}]
        )
        ans = response.choices[0].message.content
        if not ans:
            raise Exception('未知错误 无法生成文件描述')
        return ans
    except Exception as e:
        logger.exception(f"LLM API call failed: {e}")
        raise e


def format_data(local_path: str) -> str:
    """
    解析文件内容为文本：支持文档、代码、视频元数据。
    """
    path = Path(local_path)
    if not path.exists(): return ""
    ext = path.suffix.lower()

    # 1. 处理视频元数据
    video_exts = {'.mp4', '.avi', '.mov', '.mkv', '.flv', '.wmv'}
    if ext in video_exts:
        try:
            cap = cv2.VideoCapture(local_path)
            w, h = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)), int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            fps = cap.get(cv2.CAP_PROP_FPS)
            count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = count / fps if fps > 0 else 0
            cap.release()
            return f"Video Info: {path.name}, Resolution: {w}x{h}, Duration: {duration:.2f}s, FPS: {fps:.2f}"
        except:
            return f"Video File: {path.name}"

    # 2. 处理 Word 文档
    if ext == '.docx':
        try:
            doc = Document(local_path)
            content = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    content.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
            return "\n".join(content)
        except:
            return ""

    # 3. 处理文本/代码文件
    if ext in text_exts or not ext:
        try:
            return path.read_text(encoding='utf-8', errors='ignore')
        except:
            return ""

    return ""


def upload_file(local_path: str) -> str:
    """
    上传临时文件到外部服务以获取 URL。
    """
    url = "https://tmper.app/upload/"
    path = Path(local_path)
    try:
        with path.open("rb") as f:
            resp = requests.post(url, files={"file": (path.name, f)}, timeout=30)
            resp.raise_for_status()
        return resp.json()["url"]
    except Exception:
        logger.exception("upload error")
        return ""
