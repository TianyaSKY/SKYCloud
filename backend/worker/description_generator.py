"""
文件描述生成模块

使用多模态模型（VL）为文件生成智能描述：
- 支持文档、图片、视频等多种格式
- 将文件转换为图片后调用 VL 模型分析
- 生成结构化的英文描述用于后续向量化
"""

import logging
import tempfile
import base64
import mimetypes
from pathlib import Path

import cv2
from docx import Document
from openai import OpenAI

from .format_converter import (
    convert_office_to_pdf,
    convert_pdf_to_images,
    extract_video_frames,
)

# 配置日志
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# 支持的文本扩展名
TEXT_EXTENSIONS = {
    ".md",
    ".cpp",
    ".h",
    ".c",
    ".py",
    ".txt",
    ".java",
    ".go",
    ".rs",
    ".js",
    ".ts",
    ".html",
    ".css",
    ".json",
    ".yaml",
    ".yml",
    ".sql",
    ".sh",
    ".bat",
}

# 支持的文件格式分类
DOCUMENT_EXTENSIONS = {".docx", ".doc", ".xlsx", ".xls", ".pptx", ".ppt", ".pdf"}
VIDEO_EXTENSIONS = {".mp4", ".avi", ".mov", ".mkv", ".flv", ".wmv"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".flac", ".ogg"}


def image_to_base64(image_path: str) -> str:
    """
    将图片文件转换为 Base64 Data URI 字符串。
    """
    mime_type, _ = mimetypes.guess_type(image_path)
    if not mime_type:
        mime_type = "image/jpeg"

    with open(image_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode("utf-8")

    return f"data:{mime_type};base64,{encoded_string}"


def _get_visual_urls(local_path: str) -> list:
    """
    将不同类型的文件转换为可供多模态模型识别的图片 Base64 Data URI 列表。
    """
    path = Path(local_path)
    ext = path.suffix.lower()
    image_uris = []

    if ext in AUDIO_EXTENSIONS:
        raise ValueError("暂不支持音频文件的解析！！！")

    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            if ext in DOCUMENT_EXTENSIONS:
                target_pdf = (
                    local_path
                    if ext == ".pdf"
                    else convert_office_to_pdf(local_path, tmpdir)
                )
                images = convert_pdf_to_images(target_pdf, tmpdir, max_pages=20)
                for img_path in images:
                    uri = image_to_base64(img_path)
                    if uri:
                        image_uris.append(uri)

            elif ext in VIDEO_EXTENSIONS:
                frames = extract_video_frames(local_path, tmpdir, frame_count=20)
                for img_path in frames:
                    uri = image_to_base64(img_path)
                    if uri:
                        image_uris.append(uri)
            elif ext in IMAGE_EXTENSIONS:
                # 文件直接转 Base64
                image_uris.append(image_to_base64(local_path))
            elif ext in TEXT_EXTENSIONS:
                return []  # 无需处理
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        except Exception as e:
            logger.error(f"Error processing visual content for {local_path}: {e}")
            raise e

    return image_uris


def generate_file_description(local_path: str, config: dict) -> str:
    """
    生成文件的多模态描述。

    Args:
        local_path: 文件本地路径
        config: VL 模型配置，包含 api、key、model

    Returns:
        str: 生成的英文描述文本
    """
    api = config.get("api")
    key = config.get("key")
    model = config.get("model")

    # 获取图片的 Base64 Data URI 列表
    visual_contents = _get_visual_urls(local_path)
    if not visual_contents:
        logger.info(f"No visual content for {local_path}")

    client = OpenAI(api_key=key, base_url=api, timeout=120)

    # 强制要求最终输出为英文
    prompt = (
        "You are a professional file analysis assistant. Please carefully observe the provided file content "
        "(which may be document screenshots, video keyframes, or images) and generate an accurate, professional, "
        "and concise description. The description should include: core theme, main content summary, and key information points. "
        "Output the description text directly in ENGLISH without any preamble or explanatory text. Do not use Markdown syntax. "
        "No need to describe the sequence of images."
    )

    content = [{"type": "text", "text": prompt}]
    for uri in visual_contents:
        content.append({"type": "image_url", "image_url": {"url": uri}})
    content.append(
        {"type": "text", "text": _extract_text_content(local_path)}
    )  # 加入文本信息

    try:
        response = client.chat.completions.create(
            model=model, messages=[{"role": "user", "content": content}]
        )
        ans = response.choices[0].message.content
        if not ans:
            raise Exception("未知错误 无法生成文件描述")
        return ans
    except Exception as e:
        logger.exception(f"LLM API call failed: {e}")
        raise e


def _extract_text_content(local_path: str) -> str:
    """
    解析文件内容为文本：支持文档、代码、视频元数据。
    """
    path = Path(local_path)
    if not path.exists():
        return ""
    ext = path.suffix.lower()

    # 1. 处理视频元数据
    if ext in VIDEO_EXTENSIONS:
        try:
            cap = cv2.VideoCapture(local_path)
            w, h = (
                int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
            )
            fps = cap.get(cv2.CAP_PROP_FPS)
            count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = count / fps if fps > 0 else 0
            cap.release()
            return f"Video Info: {path.name}, Resolution: {w}x{h}, Duration: {duration:.2f}s, FPS: {fps:.2f}"
        except:
            return f"Video File: {path.name}"

    # 2. 处理 Word 文档
    if ext == ".docx":
        try:
            doc = Document(local_path)
            content = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    content.append(
                        " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    )
            return "\n".join(content)
        except:
            return ""

    # 3. 处理文本/代码文件
    if ext in TEXT_EXTENSIONS or not ext:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except:
            return ""

    return ""


# 为了向后兼容，保留旧函数名的别名
desc_file = generate_file_description
format_data = _extract_text_content
# upload_file 已被废弃
