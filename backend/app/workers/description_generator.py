"""文件描述生成（Worker 适配）：多模态 VL / Chat 产出中文描述供 embedding。

业务边界：只负责「文件 → 描述文本」；索引落库与向量写入在 indexing_handler。
LLM 调用统一走 llm_client.chat_completion，便于记 Token。
"""

import base64
import logging
import mimetypes
import tempfile
from pathlib import Path

import cv2
from docx import Document

from .format_converter import (
    convert_office_to_pdf,
    convert_pdf_to_images,
    extract_video_frames,
)

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# 走纯文本 Chat 路径的扩展名（比 VL 更便宜）
TEXT_EXTENSIONS = {
    ".md",
    ".cpp",
    ".h",
    ".c",
    ".py",
    ".txt",
    ".java",
    ".go",
    ".rs",
    ".js",
    ".ts",
    ".html",
    ".css",
    ".json",
    ".yaml",
    ".yml",
    ".sql",
    ".sh",
    ".bat",
}

# 走 VL 视觉路径的格式分类
DOCUMENT_EXTENSIONS = {".docx", ".doc",
                       ".xlsx", ".xls", ".pptx", ".ppt", ".pdf"}
VIDEO_EXTENSIONS = {".mp4", ".avi", ".mov", ".mkv", ".flv", ".wmv"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".flac", ".ogg"}


def image_to_base64(image_path: str) -> str:
    """图片文件 → Base64 Data URI，供 VL image_url 字段。"""
    mime_type, _ = mimetypes.guess_type(image_path)
    if not mime_type:
        mime_type = "image/jpeg"

    with open(image_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode("utf-8")

    return f"data:{mime_type};base64,{encoded_string}"


def _get_visual_urls(local_path: str) -> list:
    """按类型转成 VL 可读的图片 Data URI 列表；音频直接拒绝。"""
    path = Path(local_path)
    ext = path.suffix.lower()
    image_uris = []

    if ext in AUDIO_EXTENSIONS:
        raise ValueError("暂不支持音频文件的解析！！！")

    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            if ext in DOCUMENT_EXTENSIONS:
                target_pdf = (
                    local_path
                    if ext == ".pdf"
                    else convert_office_to_pdf(local_path, tmpdir)
                )
                images = convert_pdf_to_images(
                    target_pdf, tmpdir, max_pages=20)
                for img_path in images:
                    uri = image_to_base64(img_path)
                    if uri:
                        image_uris.append(uri)

            elif ext in VIDEO_EXTENSIONS:
                frames = extract_video_frames(
                    local_path, tmpdir, frame_count=20)
                for img_path in frames:
                    uri = image_to_base64(img_path)
                    if uri:
                        image_uris.append(uri)
            elif ext in IMAGE_EXTENSIONS:
                image_uris.append(image_to_base64(local_path))
            elif ext in TEXT_EXTENSIONS:
                return []  # 文本路径不需要视觉帧
            else:
                raise ValueError(f"Unsupported file type: {ext}")
        except Exception as e:
            logger.error(
                f"Error processing visual content for {local_path}: {e}")
            raise e

    return image_uris


def _generate_text_description(local_path: str, config: dict, user_id: int = 0) -> str:
    """纯文本文件用 Chat 模型生成描述（比 VL 更快更省 Token）。"""
    from app.services.llm_client import chat_completion

    text_content = _extract_text_content(local_path)
    if not text_content.strip():
        return "空文件"

    # 截断避免撑爆上下文
    max_chars = 8000
    if len(text_content) > max_chars:
        text_content = text_content[:max_chars] + "\n...(内容已截断)"

    prompt = (
        "你是一个专业的文件分析助手。请根据以下文本内容，生成准确、专业、简洁的描述。"
        "描述应包含：核心主题、主要内容摘要和关键信息点。"
        "请直接用中文输出描述文本，不要添加任何前缀或解释性文字。不要使用 Markdown 语法。"
    )

    try:
        response = chat_completion(
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": text_content},
            ],
            config=config,
            user_id=user_id,
            action="describe_text",
            query_summary=Path(local_path).name,
        )
        ans = response.choices[0].message.content
        if not ans:
            raise Exception("未知错误 无法生成文件描述")
        logger.info(f"[Chat] Text description generated for {local_path}")
        return ans
    except Exception as e:
        logger.exception(f"Chat LLM API call failed: {e}")
        raise e


def generate_file_description(
        local_path: str, config: dict, chat_config: dict | None = None, user_id: int = 0
) -> str:
    """生成中文文件描述：文本走 Chat，其余走 VL。

    :param local_path: 本地绝对路径
    :param config: VL 模型配置（api/key/model）
    :param chat_config: 可选 Chat 配置；缺省回退到 config
    :param user_id: 用于 Token 记账
    """
    from app.services.llm_client import chat_completion

    path = Path(local_path)
    ext = path.suffix.lower()

    if ext in TEXT_EXTENSIONS:
        text_config = chat_config or config
        return _generate_text_description(local_path, text_config, user_id=user_id)

    visual_contents = _get_visual_urls(local_path)
    if not visual_contents:
        logger.info(f"No visual content for {local_path}")

    prompt = (
        "你是一个专业的文件分析助手。请仔细观察提供的文件内容"
        "（可能是文档截图、视频关键帧或图片），生成准确、专业、简洁的描述。"
        "描述应包含：核心主题、主要内容摘要和关键信息点。"
        "请直接用中文输出描述文本，不要添加任何前缀或解释性文字。不要使用 Markdown 语法。"
        "不需要描述图片的顺序。"
    )

    content: list[dict] = [{"type": "text", "text": prompt}]
    for uri in visual_contents:
        content.append({"type": "image_url", "image_url": {"url": uri}})
    # 附带可抽取文本，补 VL 对表格/代码的盲区
    content.append(
        {"type": "text", "text": _extract_text_content(local_path)}
    )

    try:
        response = chat_completion(
            messages=[{"role": "user", "content": content}],
            config=config,
            user_id=user_id,
            action="describe_vl",
            query_summary=Path(local_path).name,
        )
        ans = response.choices[0].message.content
        if not ans:
            raise Exception("未知错误 无法生成文件描述")
        return ans
    except Exception as e:
        logger.exception(f"VL LLM API call failed: {e}")
        raise e


def _extract_text_content(local_path: str) -> str:
    """尽量抽取可索引文本：视频元数据 / docx / 纯文本代码。"""
    path = Path(local_path)
    if not path.exists():
        return ""
    ext = path.suffix.lower()

    if ext in VIDEO_EXTENSIONS:
        try:
            cap = cv2.VideoCapture(local_path)
            w, h = (
                int(cap.get(cv2.CAP_PROP_FRAME_WIDTH)),
                int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT)),
            )
            fps = cap.get(cv2.CAP_PROP_FPS)
            count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = count / fps if fps > 0 else 0
            cap.release()
            return f"Video Info: {path.name}, Resolution: {w}x{h}, Duration: {duration:.2f}s, FPS: {fps:.2f}"
        except:
            return f"Video File: {path.name}"

    if ext == ".docx":
        try:
            doc = Document(local_path)
            content = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    content.append(
                        " | ".join(c.text.strip()
                                   for c in row.cells if c.text.strip())
                    )
            return "\n".join(content)
        except:
            return ""

    if ext in TEXT_EXTENSIONS or not ext:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except:
            return ""

    return ""


# 兼容旧 import 名
desc_file = generate_file_description
format_data = _extract_text_content
